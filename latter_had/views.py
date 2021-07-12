import json
import os
from base64 import encode, decode
from datetime import date

import pdfkit as pdfkit
from django.core.files import File
from django.core.files.base import ContentFile
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt


from languages.models import Biodata, LanguageName
from .models import *
from django.http import JsonResponse
from django.http import Http404,HttpResponse
from django.template import loader

from django.shortcuts import get_object_or_404, render,redirect
from django.utils import timezone
from django import forms
from django.contrib.auth.models import Group,User
from cities_light.models import City, Country, SubRegion,Region
from django.contrib.auth import authenticate, login as dj_login, logout
from wkhtmltopdf.views import PDFTemplateResponse
import time
import os.path
from os import path

def LatterhadForm(request):
    response = {}
    id=''
    template = loader.get_template("home/latterhad_cards/latterhad_form.html")

    labels = LatterHad.objects.get(id=1)
    labels = labels.label_name
    d = json.loads(labels)
    print(type(d))
    response['labels'] = d

    try:id = request.GET['template_id']
    except:id='1'
    try:
        latterhad_card_id = request.GET['latterhad_card_id']
        response['latterhad_card_id'] = latterhad_card_id
        if latterhad_card_id != "" or latterhad_card_id != None:
            print('wed:',latterhad_card_id)
            latterhad_card = LatterHadCard.objects.get(id=latterhad_card_id)
            response['details'] = latterhad_card
            labels = LatterHad.objects.get(language=latterhad_card.language.language.id)
            labels = labels.label_name
            d = json.loads(labels)
            print(d)
            response['labels'] =d
            id = latterhad_card.template_id
            print('template_id:',id)
    except:response['latterhad_card_id'] = "-1"
    if request.user.is_authenticated:
        print('latterhad_form')
        response['template_id'] = id
        response['latterhad_template'] = LatterHadTemplateData.objects.all()
        response['all_country'] = Country.objects.all()
        languages_name = LatterHad.objects.filter(status=0)
        response['all_languages'] = LatterHad.objects.all()
        print('lang:',response['all_languages'])
        if request.user.is_authenticated:
            user = UserDetails.objects.get(user=request.user)
            user_country = user.country_dialcode
            user_country_state = Region.objects.filter(country=user_country)
            user_country_city = SubRegion.objects.filter(country=user_country)
            response['user_country'] = user_country
            response['user_country_state'] = [ a.name for a in user_country_state]
            response['user_country_city'] = [a.name for a in user_country_city]
            response['all_state'] = Region.objects.all()
        else:
            response['all_state'] = Region.objects.all()
            response['all_city'] = SubRegion.objects.all()

        return HttpResponse(template.render(response, request))
    else:
        return redirect('users:home')

        response['error'] = 'Error'
        return HttpResponse(template.render(response, request))

@csrf_exempt
def CreateLatterhadCard(request):
    function_details  = {}
    response = {}
    if request.user.is_authenticated:
        print('create')
        print(request.POST)
        data = request.POST
        # update_data_latterhad_card_id = request.GET['latterhad_card_id']
        template_id = data['template_id']
        template_instance = LatterHadTemplateData.objects.get(id=template_id)
        template_path =template_instance.template_url_latterhad

        latterhad_card_id = data['latterhad_card_id']
        language = data['language_id']

        language_name_id = language
        for file in request.FILES:
            # print(request.FILES['image'])
            try:image = request.FILES['image']
            except: image = ""
        print('------------------------------------------------:',request.user)
        user_instance = UserDetails.objects.get(user=request.user)
        latterhad_card_filter = LatterHadCard.objects.filter(id=latterhad_card_id)
        latterhad_card_exist = False
        for wd_card in latterhad_card_filter:
            latterhad_card_user = wd_card.latterhad_user
            if str(wd_card.id) == latterhad_card_id:
                print('exist',latterhad_card_id,'==',wd_card.id)
                latterhad_card_exist = LatterHadCard.objects.get(id=latterhad_card_id)

            else:
                print('card not exist')
        try:
            language_instance = LatterHad.objects.get(id=language_name_id)
        except:
            language_instance = LatterHad.objects.filter(language='english').last()
        # print(wd_card.id,':',latterhad_card_id)
        # time.sleep(20)
        if latterhad_card_exist:
            print('latterhad_card_exist')
            latterhad_card_exist.template = template_instance
            for key, values in data.items():
                try:setattr(latterhad_card_exist,key,values)
                except:print('data not exist')
            latterhad_card_exist.save()
            if request.FILES:
                try:
                    exist_image = 'media/'+str(latterhad_card_exist.image)


                    if image is not None and image != "":
                        latterhad_card_exist.image = image
                        try:
                            if str(path.isfile(exist_image)) == True:
                                os.remove(exist_image)
                        except:pass
                    latterhad_card_exist.save()
                except:
                    print('image not removed')
            latterhad_card_instance = get_latterhad_card(request, latterhad_card_id)
            response={'latterhad_card_instance':latterhad_card_instance}
            print("latterhad_card_instance:------------",latterhad_card_instance)
            print('gb_image',request.FILES)
            try:
                btemplate = loader.get_template(template_path)
            except:
                btemplate = loader.get_template('home/latterhad_cards/html_templates/latterhad_template_1.html')
            # print('btemplate',btemplate)

            html = btemplate.render(latterhad_card_instance)
        else:
            print('latterhad_card not exist')
            new_latterhad_card = LatterHadCard.objects.create(latterhad_user=user_instance,language=language_instance,template=template_instance)
            for key, values in data.items():
                try:setattr(new_latterhad_card,key,values)
                except:print('data not exist')
            if request.FILES:
                try:
                    if image is not None or image != "":
                        new_latterhad_card.image = image
                except:
                    print('image not saved')
                    time.sleep(10)
            print('card save')
            new_latterhad_card.save()
            latterhad_card_id = new_latterhad_card.id
            try:
                btemplate = loader.get_template(template_path)
            except:
                btemplate = loader.get_template('home/latterhad_cards/html_templates/latterhad_template_1.html')
            latterhad_card_instance = get_latterhad_card(request,latterhad_card_id)
            html = btemplate.render(latterhad_card_instance)
        # template = get_template(btemplate)
        print('instance',latterhad_card_instance)

        return JsonResponse({"latterhad_card_id": latterhad_card_id, "html": html}, status=200)
    else:
        return redirect('users:home')

def get_wk_pdf(request):
    print(request)
    latterhad_card_instance = {}
    show_content =True
    try:
        content_type = request.GET['content_type']
        if content_type == 'download':
            show_content = False
        else:
            show_content = True

    except:
        pass
    latterhad_card_id = request.GET['latterhad_card_id']
    latterhad_card_instance = get_latterhad_card(request,latterhad_card_id)
    filename = latterhad_card_instance['groom_first_name']+' '+latterhad_card_instance['bride_first_name']
    print('filename:',filename)
    print('-------------')

    response = PDFTemplateResponse(request=request,
                                   template=latterhad_card_instance['btemplate'],
                                   filename=filename,
                                   context=latterhad_card_instance,
                                   show_content_in_browser=show_content,
                                   cmd_options={'margin-top': 0,'margin-bottom': 0,'margin-right': 0,'margin-left': 0, },

                                   )
    # pdf = response.rendered_content
    return response
def get_latterhad_card(requset,latterhad_card_id):
    latterhad_card = LatterHadCard.objects.get(id=latterhad_card_id)
    latterhad_card_instance = LatterHadCard.objects.filter(id=latterhad_card_id).values()[0]
    template_path = LatterHadTemplateData.objects.get(id=latterhad_card_instance['template_id'])
    language_id = LanguageName.objects.get(id=latterhad_card_instance['language_id'])

    labels = LatterHad.objects.get(id=language_id.id)
    labels = labels.label_name
    template_language = json.loads(labels)
    btemplate = loader.get_template(template_path.template_url_latterhad)

    latterhad_card_instance['btemplate'] = btemplate
    latterhad_card_instance['latterhad_card_instance'] = latterhad_card_instance
    # day_name = latterhad_card_instance['date_of_latterhad'].strftime("%A")
    # day = latterhad_card_instance['date_of_latterhad'].strftime("%d")
    # month = latterhad_card_instance['date_of_latterhad'].strftime("%B")
    latterhad_card_instance['template_language'] = template_language
    print('image',latterhad_card.template.template_image_latterhad)
    a = latterhad_card_instance.update({'b_image':latterhad_card.template.template_image_latterhad})
    # print('image',month,day)
    # latterhad_card_instance['latterhad_year'] = latterhad_card_instance['date_of_latterhad'].strftime("%Y")
    # latterhad_card_instance['latterhad_day'] = day
    # latterhad_card_instance['latterhad_day_name'] = day_name
    # latterhad_card_instance['latterhad_month'] = month
    # print('image',template_language[month])
    # print(template_language)
    return latterhad_card_instance
def selected_country(request):
    if request.user.is_authenticated:
        response = {}
        data = request.POST
        try:
            country = data['country']
            country = Country.objects.get(name=country)
            country_state = Region.objects.filter(country=country)
            response['country_state'] = [a.name for a in country_state]
        except:
            country = None

        print(response)
        return JsonResponse(response, status=202)

def selected_state_city(request):
    print(request)
    if request.user.is_authenticated:
        response = {}
        data = request.POST
        try:
            state = data['state']
            state = Region.objects.get(name=state)
            state_city = SubRegion.objects.filter(region=state)
            print('state_city:-',state_city)
            response['state_city'] = [a.name for a in state_city]
            print(response)
        except:
            state = None
        print('state:',state)


        return JsonResponse(response, status=202)

def selected_language(request):
    print(request)
    if request.user.is_authenticated:
        response = {}
        data = request.POST
        print('r:',request)
        try:
            latterhad_language_id = data['latterhad_language_id']
            print('latterhad_language_id:',latterhad_language_id)
            labels = LatterHad.objects.get(id=latterhad_language_id)
            labels = labels.label_name
            d = json.loads(labels)
            print("d:-------",d)
            response['labels'] = labels
        except:
            print('error occures')
            labels = None
        print('response:',response)

        return JsonResponse(response, status=202)
def latterhadCardUpdateForm(request):
    response = {}
    data = request.GET
    latterhad_card_id = data['latterhad_card_id']
    if request.user.is_authenticated:
        print('latterhad_card_update_form')
        template = loader.get_template("home/latterhad_cards/latterhad_form.html")
        latterhad_card = LatterHadCard.objects.get(id=latterhad_card_id)
        print('biodata:',latterhad_card.language)
        response['template_id'] = latterhad_card.template_id
        response['biodata_template'] = LatterHadTemplateData.objects.all()
        response['latterhad_card_details'] = latterhad_card
        response['all_country'] = Country.objects.all()
        languages_name = LatterHad.objects.filter(status=0)
        response['all_languages'] = [{'language_name':l.language.language_name,'id':l.id} for l in languages_name]
        print('lang:',response['all_languages'])
        print(latterhad_card.language)
        labels = LatterHad.objects.get(language=latterhad_card.language.language.id)
        labels = labels.label_name
        d = json.loads(labels)
        print(d)
        response['labels'] =d
        if request.user.is_authenticated:
            user = UserDetails.objects.get(user=request.user)
            user_country = user.country_dialcode
            user_country_state = Region.objects.filter(country=user_country)
            user_country_city = SubRegion.objects.filter(country=user_country)
            response['user_country'] = user_country
            response['user_country_state'] = [ a.name for a in user_country_state]
            response['user_country_city'] = [a.name for a in user_country_city]
            response['all_state'] = Region.objects.all()
            # print('all states:',response['all_state'])
        else:
            response['all_state'] = Region.objects.all()
            response['all_city'] = SubRegion.objects.all()
        for language in response['all_languages']:
            print(language['language_name'])


        return HttpResponse(template.render(response, request))
    else:
        return redirect('users:home')

def latterhad_home(request):
    print('biodata home')
    response = {}
    template = loader.get_template("home/latterhad_cards/latterhad_home.html")
    if request.user.is_authenticated:
        user = UserDetails.objects.get(user=request.user)
        response['latterhad_drafts'] = LatterHadCard.objects.filter(latterhad_user=user,latterhad_card_status=1).order_by('created').reverse()
        response['latterhad_purchased'] = LatterHadCard.objects.filter(latterhad_user=user,latterhad_card_status=2).order_by('created').reverse()
    response['latterhad_templates'] = LatterHadTemplateData.objects.all()
    print(response)
    return HttpResponse(template.render(response, request))

@csrf_exempt
def RetrofitExample(request):
    print('retrofit')
    print(request.method)
    if request.method == "POST":
        print("post_data:",request.POST)
        print("post_data:",request.GET)
        return JsonResponse({"name":"pappu"}, safe=False)
    response = {}
    MoviesResponse = { "message":"hello","movies_layout":"abcd", "title":"baburaoganpat", "subtitle":"bach gya bechara", "description":"real based story", "rating":"5",
                       "poster_path":"poster_path", "adult":"adult", "overview":"overview", "release_date":"20/09/2020",
                       "genre_ids":"2", "id":"1", "original_title":"original_title", "original_language":"hindi",
                       "backdrop_path":"backdrop_path", "popularity":"popularity", "vote_count":"800", "video":"video", "vote_average":"58"}

    # return HttpResponse(response, status=200, content_type="application/json")
    response =[{"name":"Captain America","realname":"Steve Rogers","team":"Avengers","firstappearance":"1941","createdby":"Joe Simon","publisher":"Marvel Comics","imageurl":"https:\/\/www.simplifiedcoding.net\/demos\/marvel\/captainamerica.jpg","bio":""}]
    return JsonResponse(response, safe=False)
def CreateLanguageLabel(request):
    print('create labels for languages')
    data = request.POST
    context = {}
    label_name = {}
    try:
        language_id = request.GET['language_id']
        label_name = LatterHad.objects.values('label_name').get(id=language_id)
        print(label_name)
    except:pass

    print('request_post',request.POST)
    template = loader.get_template("home/latterhad_cards/add_language_label.html")
    context['country'] = Country.objects.all()

    if request.method == 'POST':
        language_name = data['language_name']
        country_name = data['country']
        language_abr = data['language_abr']
        for key, values in data.items():
            if key != 'csrfmiddlewaretoken' and key != 'language_name' and key != 'language_abr' and key != 'country':
                label_name[key] = values
        a = {}
        label_name=json.dumps(label_name)

        print('label:',label_name)

        language_existance = LanguageName.objects.filter(language_name__icontains=language_name).last()

        if language_existance:
            label_instance = LatterHad.objects.filter(language=language_existance)
            if label_instance.exists():
                print('language exists')
            else:
                add_language_label = LatterHad.objects.create(language=language_existance, label_name=label_name,status=True)
        else:
            print('language not exists')
            add_language = LanguageName.objects.create(language_name=language_name, country_name=country_name,language_abr=language_abr)
            add_language_label = LatterHad.objects.create(language=add_language,label_name=label_name,status=True)
    return HttpResponse(template.render(context,request))

from docx import *
from docx.shared import Inches

def word_docs(request):
    document = Document()
    data = request.GET
    response = {}
    docx_title=str('letterhead')+".docx"
    # ---- Cover Letter ----
    latterhad = LatterHadTemplateData.objects.all().last()

    if request.method == 'GET':
        id = data['latterhad_card_id']
        latterhad_instance = LatterHadCard.objects.get(id=id)
        docx_title=str(latterhad_instance.company_name)+".docx"
        header_section = document.sections[0]


        header = header_section.header

        paragraph = header.paragraphs[0]

        paragraph.text = str(latterhad.template_image_latterhad)
        paragraph.style = document.styles["Header"]

        import xlsxwriter


        workbook = xlsxwriter.Workbook('headers_footers.xlsx')
        worksheet = workbook.add_worksheet('Image')

        # Adjust the page top margin to allow space for the header image.
        worksheet.set_margins(top=1.3)

        # worksheet.set_header('&L&[Picture]&C&16&"Calibri,Bold"Revenue Report',
        #                      {'image_left': latterhad.template_image_latterhad})

        workbook.close()
        document.sections[0].header.add_picture(latterhad.template_image_latterhad, width=Inches(4))



    document.add_picture(latterhad.template_image_latterhad, width=Inches(4))
    document.add_heading
    # document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))
    # document.add_paragraph('Dear Sir or Madam:')
    # document.add_page_break()
    # Prepare document for download

    # -----------------------------
    from htmldocx import HtmlToDocx

    new_parser = HtmlToDocx()
    # btemplate = loader.get_template("home/latterhad_cards/html_templates/latterhad_template_1.html")
    btemplate = "/home/hira/PycharmProjects/CardsMaker/templates/home/latterhad_cards/html_templates/latterhad_template_1.html"
    print('template:',btemplate)
    k = new_parser.parse_html_file(btemplate, "htmltodocs.docs")
    from io import BytesIO
    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)

    response = HttpResponse(
        # f.getvalue(),
        f,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    print(k)
    return response